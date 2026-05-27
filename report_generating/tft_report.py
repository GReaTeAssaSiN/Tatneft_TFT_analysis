"""
Формирование DOCX-отчёта по TFT-модели.
Описывает источники данных, таксономию переменных, предобработку и архитектуру TFT-модели.

Запускать из корня проекта: python report_generating/tft_report.py
Результат: reports/tft_report.docx
"""

import os
import sys

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from utils.data_utils import (
    EXCLUDED_COLS,
    LOG_COLS,
    NO_ZSCORE_COLS,
    STATIC_CATS,
    STATIC_REALS,
    TARGET_COLS,
    TIME_VARYING_KNOWN_CATS,
    TIME_VARYING_KNOWN_REALS,
    TIME_VARYING_UNKNOWN_REALS,
    ENCODER_LENGTH,
    PREDICTION_LENGTH,
)

os.makedirs("reports", exist_ok=True)

# ── Проверка наличия данных ──────────────────────────────────
_src_prepared = "data/prepared_data.csv"
_src_merged   = "data/merged_data.csv"
_src_static   = "SourceDataForWork/gas_stations_static.csv"
_src_temporal = "SourceDataForWork/gas_stations_temporal_daily_2025.csv"

if not os.path.exists(_src_prepared) or not os.path.exists(_src_merged):
    missing = [p for p in (_src_prepared, _src_merged) if not os.path.exists(p)]
    print(f"[ОШИБКА] Файлы не найдены: {missing}")
    print("  Запустите: python explore_data.py && python eda/eda_preprocessing.py")
    sys.exit(1)

df     = pd.read_csv(_src_prepared, parse_dates=["date"])
merged = pd.read_csv(_src_merged,   parse_dates=["date"])

# Точное число колонок исходных файлов (загружаем только заголовок)
if os.path.exists(_src_static):
    _static_cols = len(pd.read_csv(_src_static, encoding="utf-8-sig", nrows=0).columns)
else:
    _static_cols = "—"
if os.path.exists(_src_temporal):
    _temporal_cols = len(pd.read_csv(_src_temporal, encoding="utf-8-sig", nrows=0).columns)
else:
    _temporal_cols = "—"

# ── Вспомогательные функции ──────────────────────────────────
def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x37, 0x86, 0x8C)
    return p


def add_table(doc, headers, rows, col_widths=None, header_color="1F3964"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        cell_bg(hdr[i], header_color)
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cell_bg(cells[ci], bg)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(9)

    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return tbl


def para(doc, text, bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


# ============================================================
# Документ
# ============================================================
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2)

# ── Титул ────────────────────────────────────────────────────
title = doc.add_heading("TFT-модель прогнозирования продаж АЗС", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    run.font.size = Pt(18)

subtitle = doc.add_paragraph("Описание данных, переменных, предобработки и архитектуры модели")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph(
    f"Temporal Fusion Transformer (Lim et al., 2020)  |  "
    f"Период: 2025 год  |  "
    f"Станций: {merged['station_id'].nunique()}  |  "
    f"Гранулярность: ежедневные данные"
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================================
# 1. ИСТОЧНИКИ ДАННЫХ
# ============================================================
h1(doc, "1. Источники данных")

para(doc, "В работе используются два исходных файла для 5 АЗС Татнефть:", bold=True)

add_table(
    doc,
    headers=["Файл", "Строк", "Колонок", "Содержимое"],
    rows=[
        ["gas_stations_static.csv",
         str(merged["station_id"].nunique()), str(_static_cols),
         "Паспорт АЗС: тип дороги, услуги, количество колонок, характеристики расположения"],
        ["gas_stations_temporal_daily_2025.csv",
         str(merged["station_id"].nunique() * 365), str(_temporal_cols),
         "Ежедневные данные 2025 г.: погода, трафик, продажи, цены, акции, метрики клиентов"],
        [f"merged_data.csv (LEFT JOIN по station_id)",
         str(merged.shape[0]), str(merged.shape[1]),
         "Объединённый датасет: все временные данные + статические характеристики АЗС"],
        [f"prepared_data.csv (после предобработки)",
         str(df.shape[0]), str(df.shape[1]),
         "Нормализованный датасет для подачи в TimeSeriesDataSet"],
    ],
    col_widths=[5.5, 2, 2.5, 8],
)

doc.add_paragraph()
para(
    doc,
    f"LEFT JOIN добавляет к каждой из {merged.shape[0]} строк временного ряда "
    f"статические признаки из паспорта АЗС. "
    f"Временной ряд: {merged['date'].min().date()} — {merged['date'].max().date()} "
    f"({merged['date'].nunique()} дней x {merged['station_id'].nunique()} станций). "
    "Пропущенных дат нет.",
)

# ============================================================
# 2. ИСКЛЮЧЁННЫЕ ПЕРЕМЕННЫЕ
# ============================================================
h1(doc, "2. Исключённые переменные (не подаются в модель)")

n_merged = merged.shape[1]
n_after_excl = n_merged - len(EXCLUDED_COLS)
# preprocessing adds: _enc (7) + sin/cos (6) + time_idx (1) + _orig (12) = 26
n_added = df.shape[1] - n_after_excl
para(
    doc,
    f"Из {n_merged} колонок объединённого датасета {len(EXCLUDED_COLS)} исключены как избыточные. "
    f"Остаётся {n_merged} - {len(EXCLUDED_COLS)} = {n_after_excl} колонок. "
    f"Препроцессинг добавляет {n_added} новых колонок "
    f"(7 _enc + 6 sin/cos + 1 time_idx + 12 _orig). "
    f"Итого в prepared_data.csv: {df.shape[1]} колонок.",
)

add_table(
    doc,
    headers=["Переменная", "Причина исключения"],
    rows=[
        ["station_name",
         "Служебная строка; идентификатор станции — station_id"],
        ["total_pumps",
         "= sum(num_pumps_AI92, num_pumps_AI95, …) — дублирует детальные колонки"],
        ["total_fuel_sales",
         "= sum(sales_AI92, sales_AI95, …) — дублирует целевые переменные"],
        ["quarter",
         "Однозначно выводится из month; избыточен для модели"],
        ["traffic_uroven_jbslugi_1_poputn",
         "Уровень обслуживания дороги A–F (строка): дублирует числовые счётчики трафика"],
        ["traffic_uroven_jbslugi_2_poputn",
         "То же, полоса 2, попутное направление"],
        ["traffic_uroven_jbslugi_1_wstrechn",
         "Уровень обслуживания дороги A–F, полоса 1, встречное направление"],
        ["traffic_uroven_jbslugi_2_wstrechn",
         "То же, полоса 2, встречное направление"],
    ],
    col_widths=[5.5, 13],
)

# ============================================================
# 3. ТАКСОНОМИЯ ПЕРЕМЕННЫХ TFT
# ============================================================
h1(doc, "3. Таксономия переменных TFT")

para(
    doc,
    "Temporal Fusion Transformer (Lim et al., 2020) принимает три типа входов. "
    "Разделение определяет, какую информацию модель использует при прогнозировании: "
    "статические — зашиты в контекстный вектор; известные будущие — подаются в декодер; "
    "наблюдаемые прошлые — только в энкодер.",
    bold=True,
)

add_table(
    doc,
    headers=["Тип входа TFT", "Определение", "Кол-во"],
    rows=[
        ["static_categoricals",
         "Категориальные свойства АЗС, не меняющиеся во времени",
         str(len(STATIC_CATS))],
        ["static_reals",
         "Числовые свойства АЗС из паспорта (metadata). Не нормализуются.",
         str(len(STATIC_REALS))],
        ["time_varying_known_categoricals",
         "Категориальные признаки, известные заранее или задаваемые в what-if",
         str(len(TIME_VARYING_KNOWN_CATS))],
        ["time_varying_known_reals",
         "Числовые признаки, известные на весь горизонт прогноза "
         "(время, погода, трафик, цены). "
         "Только эти переменные можно менять в сценарном what-if анализе.",
         str(len(TIME_VARYING_KNOWN_REALS))],
        ["time_varying_unknown_reals",
         "Наблюдаемые прошлые: известны за прошлые периоды, "
         "недоступны для горизонта прогноза. Клиентские метрики и производные трафика.",
         str(len(TIME_VARYING_UNKNOWN_REALS))],
        ["target",
         "12 целевых переменных (5 видов топлива + 7 категорий магазина). "
         "TFT использует их как авторегрессивные входы энкодера автоматически.",
         str(len(TARGET_COLS))],
    ],
    col_widths=[5.5, 11, 2],
)

# ── 3.1 Статические переменные ────────────────────────────────
h2(doc, "3.1 Статические переменные")

para(
    doc,
    "Паспортные данные АЗС из gas_stations_static.csv. "
    "Не изменяются во времени. "
    "TFT строит для них глобальный статический энкодер (контекстный вектор). "
    "Категориальные не нормализуются (эмбеддинги); числовые — Z-score не применяется.",
)

add_table(
    doc,
    headers=["Переменная (оригинал)", "Тип TFT", "Предобработка", "Описание"],
    rows=[
        # STATIC CATS
        ["road_type",          "static_categoricals",
         "LabelEncoder -> road_type_enc",
         "Тип дороги (федеральная / региональная / местная)"],
        ["road_level",         "static_categoricals",
         "LabelEncoder -> road_level_enc",
         "Уровень дороги (числовой код: 1, 5, …). Порядковая категория, не число."],
        ["direction",          "static_categoricals",
         "LabelEncoder -> direction_enc",
         "Направление (транзит-трассовая / городская и др.)"],
        ["settlement_size",    "static_categoricals",
         "LabelEncoder -> settlement_size_enc",
         "Размер ближайшего населённого пункта (код 1–6). Порядковая категория."],
        ["distance_to_city_km","static_categoricals",
         "LabelEncoder -> distance_to_city_km_enc",
         "Расстояние до ближайшего города (код 0–6). Порядковая категория."],
        # STATIC REALS
        ["shop_area_m2",       "static_reals", "—",
         "Площадь магазина, м²"],
        ["num_pumps_AI92/95/AI100_bio/DT/DT_bio/... (10 к.)",
         "static_reals", "— (целые числа)",
         "Количество колонок по каждому виду топлива (AI92, AI92_bio, AI95, AI95_bio, "
         "AI100_bio, DT, DT_bio, SUG, KPG, SPG)"],
        ["has_car_wash / has_tire_service / has_cafe / has_hotel / has_shop (5 к.)",
         "static_reals", "— (бинарные)",
         "Наличие основных сервисов (0/1)"],
        ["has_shop_молельная_комната / has_shop_прачечная / "
         "has_shop_электрозарядная_станция / has_shop_подкачка_шин (4 к.)",
         "static_reals", "— (бинарные)",
         "Наличие дополнительных услуг магазина (0/1)"],
        ["competitors_wink",   "static_reals", "—",
         "Число конкурирующих АЗС поблизости (данные Wink)"],
    ],
    col_widths=[5, 3.5, 4, 6],
)

# ── 3.2 Известные будущие переменные ────────────────────────────────
h2(doc, "3.2 Известные будущие переменные")

para(
    doc,
    "Признаки, задаваемые заранее (расписание, метеопрогноз, плановые акции) "
    "или устанавливаемые пользователем в what-if сценарии. "
    "TFT подаёт их и в энкодер (ретроспектива), и в декодер (горизонт прогноза). "
    "Именно поэтому к ним применима сценарная аналитика.",
)

add_table(
    doc,
    headers=["Переменная / Группа", "Тип TFT", "Предобработка", "Обоснование / what-if"],
    rows=[
        # KNOWN CATS
        ["season",
         "known_categoricals",
         "LabelEncoder -> season_enc",
         "Сезон известен заранее; what-if: другой сезон"],
        ["holiday_name",
         "known_categoricals",
         "NaN->'нет_праздника'; LabelEncoder",
         "Государственные праздники известны заранее"],
        # KNOWN REALS — время
        ["day_of_week, day_of_week_sin, day_of_week_cos (3)",
         "known_reals", "sin/cos(2*pi*d/7)",
         "День недели; циклическое кодирование устраняет разрыв Вс->Пн"],
        ["week_of_year, week_of_year_sin, week_of_year_cos (3)",
         "known_reals", "sin/cos(2*pi*w/52)",
         "Номер недели (сезонность)"],
        ["month, month_sin, month_cos (3)",
         "known_reals", "sin/cos(2*pi*m/12)",
         "Месяц (месячная сезонность)"],
        ["is_weekend",         "known_reals", "— (бинарная)",
         "Выходной день"],
        ["is_holiday",         "known_reals", "— (бинарная)",
         "Праздничный день"],
        # KNOWN REALS — погода
        ["weather_condition",  "known_reals", "— (бинарный 0/1)",
         "Тип погоды: 0 = ясно/облачно, 1 = осадки/плохая видимость; "
         "задаётся по метеопрогнозу; what-if: 'что если осадки?'"],
        ["temperature",        "known_reals", "Z-score (только train)",
         "Температура воздуха, °C; what-if: ожидаемая температура"],
        ["precipitation_mm",   "known_reals", "Z-score (только train)",
         "Осадки, мм; what-if: интенсивность осадков"],
        # KNOWN REALS — акции
        ["promotion_fuel_active", "known_reals", "— (бинарная)",
         "Акция на топливо активна; what-if: включить/выключить"],
        ["promotion_shop_active",  "known_reals", "— (бинарная)",
         "Акция в магазине; what-if: эффект промо-кампании"],
        ["promotion_cafe_active",  "known_reals", "— (бинарная)",
         "Акция в кафе; what-if"],
        # KNOWN REALS — цены
        ["price_AI92 / AI95 / DT / DT_bio / AI100_bio (5 к.)",
         "known_reals", "Z-score (только train)",
         "Текущие цены топлива, руб/л; ежедневно меняются в 2025 г.; "
         "what-if: сценарий ценообразования"],
        # KNOWN REALS — трафик (счётчики)
        ["traffic_Passengers_cars / Truck_short / Truck / Truck_long / "
         "Transporter / Undefined, полоса 1+2, попутн. (12 к.)",
         "known_reals", "Z-score (только train)",
         "Счётчики ТС по типам, попутное направление; "
         "прогнозируются службами дорожного движения; what-if: высокий/низкий трафик"],
        ["traffic_* встречное направление, полоса 1+2 (12 к.)",
         "known_reals", "Z-score (только train)",
         "Счётчики ТС, встречное направление"],
        # KNOWN REALS — конкуренты
        ["competitor_price_AI92 / AI95 / DT (3 к.)",
         "known_reals", "Z-score (только train)",
         "Цены небрендовых конкурентов; what-if: снижение/повышение"],
        ["competitor_price_AI92_brend / AI95_brend / DT_brend / AI100 (4 к.)",
         "known_reals", "Z-score (только train)",
         "Цены брендовых конкурентов + AI100 конкурентов; what-if"],
    ],
    col_widths=[5, 3, 4, 6.5],
)

# ── 3.3 Наблюдаемые прошлые ──────────────────────────────────
h2(doc, "3.3 Наблюдаемые прошлые (time_varying_unknown_reals)")

para(
    doc,
    f"В данном проекте TIME_VARYING_UNKNOWN_REALS содержит {len(TIME_VARYING_UNKNOWN_REALS)} переменных: "
    "клиентские метрики и производные показатели трафика.",
)
para(
    doc,
    "Эти переменные ИЗВЕСТНЫ за прошлые периоды (энкодер видит их), "
    f"но НЕ МОГУТ быть заданы для горизонта прогноза ({PREDICTION_LENGTH} дн.): "
    "они не прогнозируются службами заранее и не задаются в what-if сценарии.",
)
para(
    doc,
    "Трафик-счётчики (количество ТС по типам, полоса 1 и 2) перенесены в KNOWN_REALS: "
    "прогнозируются дорожными службами. "
    "Производные метрики трафика (скорость, плотность, интенсивность) остаются в UNKNOWN: "
    "они вычисляются из счётчиков и требуют симуляции дорожного потока.",
)

add_table(
    doc,
    headers=["Переменная / Группа", "Описание"],
    rows=[
        ["corporate_customer_ratio",
         "Доля корпоративных клиентов в день. Наблюдается по факту, не прогнозируется."],
        ["customer_loyalty_score",
         "Оценка лояльности клиентов. Меняется ежедневно, не прогнозируется заранее."],
        ["traffic_scorost_1_poputn / 2_poputn / 1_wstrechn / 2_wstrechn (4 к.)",
         "Средняя скорость потока, км/ч. Производная от объёма ТС."],
        ["traffic_plotnost_1/2 (poputn+wstrechn = 4 к.)",
         "Плотность потока, авт/км."],
        ["traffic_intensiv_priv_1/2 (poputn+wstrechn = 4 к.)",
         "Интенсивность приведённого потока, авт/ч."],
        ["traffic_intensiv_fiz_1/2 (poputn+wstrechn = 4 к.)",
         "Интенсивность физического потока, авт/ч."],
    ],
    col_widths=[6, 12.5],
)

# ── 3.4 Целевые переменные ────────────────────────────────────
h2(doc, "3.4 Целевые переменные (12 штук)")

para(
    doc,
    "TFT прогнозирует 12 переменных совместно, используя взаимные корреляции "
    "(покупатели топлива = покупатели магазина). Модель выдаёт квантильные "
    "прогнозы [q2, q10, q25, q50, q75, q90, q98] для каждой цели.",
)

fuel_cols = [c for c in TARGET_COLS if c.startswith("sales_")]
shop_cols = [c for c in TARGET_COLS if c.startswith("shop_")]

skew_rows = []
for col in TARGET_COLS:
    orig_col = col + "_orig"
    if orig_col in df.columns:
        skew_before = df[orig_col].skew()
        skew_after  = df[col].skew()
        mn = df[orig_col].min()
        mx = df[orig_col].max()
        unit = "л/день" if col.startswith("sales_") else "руб/день"
        skew_rows.append([col, f"{mn:.0f}-{mx:.0f} {unit}",
                          f"{skew_before:.2f}", f"{skew_after:.2f}"])
    else:
        skew_rows.append([col, "-", "-", "-"])

add_table(
    doc,
    headers=["Переменная", "Диапазон", "Skew до log1p", "Skew после log1p"],
    rows=skew_rows,
    col_widths=[6, 4.5, 3, 3],
)

doc.add_paragraph()
para(
    doc,
    "log1p(x) = log(1+x) снижает правостороннюю асимметрию продаж (skew 2-4 -> 0.2-1.5). "
    "Работает при x=0 (дни с нулевыми продажами). "
    "TorchNormalizer(method='robust') внутри TFT дополнительно нормализует по медиане и IQR.",
)

# ============================================================
# 4. ПРЕДОБРАБОТКА — ПОШАГОВОЕ ОПИСАНИЕ
# ============================================================
h1(doc, "4. Предобработка данных")

para(doc, "Файл: eda/eda_preprocessing.py. Порядок шагов строго фиксирован.", bold=True)

# Вычислим реальные числа для описаний
n_raw    = merged.shape[1]   # 115
n_excl   = len(EXCLUDED_COLS)  # 8
n_after2 = n_raw - n_excl     # 107
n_enc_new = 7                  # 4 string + 3 numeric cats
n_sincos  = 6                  # 3 features x 2
n_after3  = n_after2 + n_enc_new    # 114
n_after4  = n_after3 + n_sincos     # 120
n_after5  = n_after4 + 1            # 121 (time_idx)
n_after6  = n_after5 + len(TARGET_COLS)  # 133 (_orig)

steps = [
    (
        "Шаг 1. Заполнение пропусков",
        "Одно поле содержит пропуски:\n"
        "  - holiday_name: заполняется значением 'нет_праздника'\n"
        "    (~99% строк — обычные дни без праздника)\n"
        "Удаление строк с пропусками недопустимо: нарушило бы непрерывность временного ряда.",
    ),
    (
        f"Шаг 2. Исключение избыточных колонок ({n_excl} штук)",
        f"Из датафрейма удаляются {n_excl} избыточных колонок (см. раздел 2):\n"
        "  - station_name  — служебная строка\n"
        "  - total_pumps   — = sum(num_pumps_*), дублирует детальные\n"
        "  - total_fuel_sales — = sum(sales_*), дублирует целевые\n"
        "  - quarter       — выводится из month\n"
        "  - traffic_uroven_jbslugi_1/2_poputn — уровень обслуживания A-F, строки\n"
        "  - traffic_uroven_jbslugi_1/2_wstrechn — то же, встречное направление\n"
        f"После исключения: {n_raw} - {n_excl} = {n_after2} колонок.\n"
        "Winsorization не применяется: не описана в статье TFT (Lim et al., 2020); "
        "выбросы целевых переменных обрабатываются TorchNormalizer(robust) внутри модели.",
    ),
    (
        "Шаг 3. Label Encoding категориальных переменных",
        f"Закодировано {n_enc_new} колонок:\n"
        "  Строковые категории (4 шт.): road_type, direction, season, holiday_name\n"
        "  Числовые коды категорий (3 шт.): road_level, settlement_size, distance_to_city_km\n"
        "    -> эти поля содержат коды (1=малый город, 6=мегаполис), не числа\n"
        "    -> нормализация Z-score некорректна; нужен LabelEncoder + эмбеддинг\n"
        "Для каждой колонки создаётся новая с суффиксом _enc.\n"
        f"Итого: {n_after2} + {n_enc_new} = {n_after3} колонок.\n"
        "LabelEncoder обучается (fit) на ВСЕХ 1825 строках, чтобы словарь\n"
        "содержал все категории, включая праздники только в ноябре/декабре.\n"
        "Для категориального кодирования (не для обучения модели) это допустимо.",
    ),
    (
        "Шаг 4. Циклическое sin/cos кодирование",
        "Ежедневные данные имеют 3 цикличных признака (hour отсутствует):\n"
        "  day_of_week (период 7):  day_of_week_sin, day_of_week_cos\n"
        "  week_of_year (период 52): week_of_year_sin, week_of_year_cos\n"
        "  month (период 12):        month_sin, month_cos\n"
        "Проекция на единичную окружность: col_sin = sin(2*pi*col/period)\n"
        "Расстояние между Вс (day=6) и Пн (day=0) на окружности минимально.\n"
        f"Добавляются 6 новых колонок. Итого: {n_after3} + {n_sincos} = {n_after4}.",
    ),
    (
        "Шаг 5. Монотонный индекс времени (time_idx)",
        "TFT требует целочисленный time_idx без пропусков для каждой группы (станции).\n"
        "Реализация: cumcount() после сортировки по (station_id, date).\n"
        f"  time_idx = 0, 1, 2, ..., 364 для каждой из {merged['station_id'].nunique()} "
        "станций (365 дней 2025 года).\n"
        f"Добавляется 1 колонка. Итого: {n_after4} + 1 = {n_after5}.",
    ),
    (
        "Шаг 6. Логарифмическое преобразование (log1p)",
        f"Все {len(TARGET_COLS)} целевых переменных (TARGET_COLS) имеют правостороннюю асимметрию.\n"
        "log1p(x) = log(1+x):\n"
        "  - работает при x=0 (нулевые продажи редких позиций)\n"
        "  - снижает skew с 2-4 до 0.2-1.5\n"
        "  - улучшает сходимость нейронной сети\n"
        "Оригинальные значения сохраняются в _orig (для обратного преобразования в predict.py).\n"
        f"Добавляются {len(TARGET_COLS)} _orig колонок. "
        f"Итого: {n_after5} + {len(TARGET_COLS)} = {n_after6}.",
    ),
    (
        "Шаг 7. Z-score нормализация per-station (без data leakage)",
        "Формула: z = (x - mean_train) / std_train\n"
        "Статистика вычисляется ТОЛЬКО по обучающей выборке (Jan-Oct 2025).\n"
        "Одни и те же mean/std применяются к val и test — стандартная ML-практика,\n"
        "исключающая data leakage из будущего в нормализацию обучающих данных.\n"
        "\n"
        "Исключения из нормализации:\n"
        "  - STATIC_REALS: паспортные данные АЗС (константны для каждой станции)\n"
        "  - Бинарные переменные (0/1): нормализация бессмысленна\n"
        "  - _enc колонки: целые числа для эмбеддинга\n"
        "  - _orig колонки: оригиналы до log1p\n"
        "  - TARGET_COLS: нормализует TorchNormalizer внутри TFT\n"
        f"  - NO_ZSCORE_COLS = {NO_ZSCORE_COLS if NO_ZSCORE_COLS else '[]'}: "
        "в данных 2025 г. цены меняются ежедневно (std > 0 везде)\n"
        "\n"
        "Скейлеры (mean, std) сохраняются в tft/scalers.pkl "
        "для обратного преобразования прогнозов в predict.py.",
    ),
]

for title_step, desc in steps:
    h3(doc, title_step)
    for line in desc.split("\n"):
        style = "List Bullet" if line.startswith("  -") else "Normal"
        p = doc.add_paragraph(
            line.lstrip("  ") if style == "List Bullet" else line,
            style=style,
        )
        for run in p.runs:
            run.font.size = Pt(10)

# ── Сплиты ────────────────────────────────────────────────────
h3(doc, "Темпоральный сплит данных")

para(
    doc,
    "Данные разбиваются строго по времени (случайная выборка недопустима — data leakage):",
)

add_table(
    doc,
    headers=["Сплит", "Период df, переданного в from_dataset()", "Строк", "Доля train", "Декодерные окна (целевой период)"],
    rows=[
        ["Train",
         "01.01.2025 - 31.10.2025",
         "1520", "~83%",
         "Январь–октябрь (все окна). Вычисляются mean/std для Z-score."],
        ["Val",
         "01.01.2025 - 30.11.2025",
         "1670", "—",
         "Ноябрь. Окна с декодером в ноябре."],
        ["Test",
         "01.01.2025 - 31.12.2025",
         "1825", "—",
         "Декабрь. Окна с декодером в декабре."],
    ],
    col_widths=[2, 5, 1.5, 1.5, 8.5],
)

doc.add_paragraph()
para(
    doc,
    f"Почему Val и Test включают весь год, а не только «свой» месяц?",
    bold=True,
)
para(
    doc,
    f"TimeSeriesDataSet.from_dataset() создаёт сэмплы для ВСЕХ допустимых окон "
    f"(encoder + decoder) внутри переданного датафрейма. "
    f"Энкодерное окно первого ноябрьского дня (1 ноября) требует {ENCODER_LENGTH} дней "
    f"ретроспективы — октябрь 1–31. "
    f"Если передать только ноябрь, первые окна получат неполный контекст (encoder_length < {ENCODER_LENGTH}), "
    f"что ухудшает качество предсказаний. "
    f"Аналогично для декабря: первые декабрьские окна требуют ноябрьского контекста. "
    f"Поэтому Val df = январь–ноябрь, Test df = весь январь–декабрь.",
    italic=True,
)
para(
    doc,
    f"Важно: «лишние» окна (энкодер=январь, декодер=январь) тоже попадают в Val/Test df, "
    f"но модель уже видела их при обучении — это нормально. "
    f"«Честными» метриками остаются только ноябрьские (val) и декабрьские (test) декодерные окна. "
    f"max_encoder_length={ENCODER_LENGTH} гарантирует, что модель не видит данные «из будущего».",
    italic=True,
)

# ============================================================
# 5. ПАРАМЕТРЫ TFT-МОДЕЛИ
# ============================================================
h1(doc, "5. Параметры TFT-модели")

add_table(
    doc,
    headers=["Параметр", "CPU", "GPU", "Описание"],
    rows=[
        ["hidden_size",           "64",  "128",
         "Размер скрытого слоя"],
        ["attention_head_size",   "2",   "4",
         "Число голов многоголового внимания (temporal self-attention)"],
        ["hidden_continuous_size","32",  "64",
         "Размер Variable Selection Network для числовых входов "
         "(рекомендация TFT: hidden_continuous <= hidden_size / 2)"],
        ["BATCH_SIZE",            "32",  "64",
         "Размер мини-батча"],
        ["EPOCHS",                "50",  "80",
         "Максимальное число эпох (EarlyStopping остановит раньше)"],
    ],
    col_widths=[5, 2, 2, 9.5],
)

doc.add_paragraph()
add_table(
    doc,
    headers=["Параметр", "Значение", "Описание"],
    rows=[
        ["ENCODER_LENGTH",
         f"{ENCODER_LENGTH} дн.",
         f"Ретроспективное окно. Охватывает месячный и недельный циклы (1 месяц)."],
        ["PREDICTION_LENGTH",
         f"{PREDICTION_LENGTH} дн.",
         f"Горизонт прогноза: следующая неделя. "
         "Поддерживает рекурсивное прогнозирование на 1 день (первый шаг) "
         "и на 1 месяц (4-5 итераций по 7 дней)."],
        ["learning_rate",     "3e-4",
         "Начальный LR. Снижается ReduceLROnPlateau(patience=5)"],
        ["dropout",           "0.15",
         "Регуляризация (Dropout + VariationalDropout)"],
        ["gradient_clip",     "1.0",
         "Обрезка градиентов (предотвращает взрыв градиентов)"],
        ["EarlyStopping",     "patience=12",
         "Останавливается при отсутствии улучшения val_loss 12 эпох подряд"],
        ["loss",
         "QuantileLoss([0.02, 0.1, 0.25, 0.5, 0.75, 0.9, 0.98])",
         "7-квантильный прогноз: медиана + доверительные интервалы"],
        ["target_normalizer",
         "MultiNormalizer([TorchNormalizer(robust)] x 12)",
         "Нормализация целей по медиане+IQR (устойчива к выбросам). "
         "Применяется внутри TFT поверх log1p."],
        ["add_relative_time_idx", "True",
         "Относительная позиция шага в окне -> known_real"],
        ["add_target_scales",     "True",
         "Масштаб каждой цели (location + scale) -> static_real"],
        ["add_encoder_length",    "True",
         "Фактическая длина энкодера -> static_real"],
        ["min_encoder_length",
         f"{ENCODER_LENGTH // 2} дн.",
         "Минимальное окно (половина от max_encoder_length)"],
    ],
    col_widths=[4.5, 5, 9],
)

# ============================================================
# 6. ИТОГОВЫЙ СОСТАВ ТЕНЗОРОВ TFT
# ============================================================
h1(doc, "6. Итоговый состав тензоров TFT")

para(
    doc,
    "Размеры тензоров, которые TFT получает на вход. "
    "Включают как исходные переменные, "
    "так и дополнительные признаки, добавляемые pytorch-forecasting автоматически.",
    bold=True,
)

# Вычисляем теоретические размеры
n_static_cats   = len(STATIC_CATS)               # 5
n_static_reals  = len(STATIC_REALS)              # 21
n_target_scales = len(TARGET_COLS) * 2           # 24 (location + scale)
n_enc_length    = 1
n_static_total  = n_static_cats + n_static_reals + n_target_scales + n_enc_length

n_known_cats    = len(TIME_VARYING_KNOWN_CATS)   # 2
n_known_reals   = len(TIME_VARYING_KNOWN_REALS)  # 53
n_rel_time      = 1
n_unknown       = len(TIME_VARYING_UNKNOWN_REALS) # 18
n_targets_enc   = len(TARGET_COLS)               # 12
n_enc_total     = n_known_cats + n_known_reals + n_rel_time + n_unknown + n_targets_enc

n_dec_total     = n_known_cats + n_known_reals + n_rel_time

h2(doc, f"6.1 Статические признаки")
add_table(
    doc,
    headers=["Компонент", "Кол-во", "Параметр датасета", "Описание"],
    rows=[
        ["static_categoricals",  str(n_static_cats),
         "static_categoricals",
         "road_type_enc, road_level_enc, direction_enc, settlement_size_enc, distance_to_city_km_enc"],
        ["static_reals (паспорт)", str(n_static_reals),
         "static_reals",
         "shop_area_m2, num_pumps_* (10), has_* (9), competitors_wink"],
        ["target_scales", str(n_target_scales),
         "add_target_scales=True",
         "12 целей x 2: location (медиана) + scale (IQR) от TorchNormalizer(robust)"],
        ["encoder_length", str(n_enc_length),
         "add_encoder_length=True",
         "Фактическая длина энкодера в текущем сэмпле"],
        ["ИТОГО", str(n_static_total), "-", ""],
    ],
    col_widths=[4.5, 2, 4.5, 7.5],
    header_color="1F3964",
)

doc.add_paragraph()
h2(doc, f"6.2 Признаки энкодера (прошлое)")
add_table(
    doc,
    headers=["Компонент", "Кол-во", "Параметр датасета", "Описание"],
    rows=[
        ["known_categoricals", str(n_known_cats),
         "time_varying_known_categoricals",
         "season_enc, holiday_name_enc"],
        ["known_reals", str(n_known_reals),
         "time_varying_known_reals",
         f"Время (9: raw+sin/cos x 3 признака), флаги (2), "
         "weather_condition (1), акции (3), "
         "цены топлива (5), погода (2), трафик поп. (12), трафик встр. (12), конкуренты (7)"],
        ["relative_time_idx", str(n_rel_time),
         "add_relative_time_idx=True",
         "Относительная позиция шага в энкодерном окне [0..1]"],
        ["unknown_reals", str(n_unknown),
         "time_varying_unknown_reals",
         "Клиентские метрики (2) + производные трафика: "
         "скорость (4), плотность (4), интенс. привед. (4), интенс. физ. (4)"],
        ["targets (observed past)", str(n_targets_enc),
         "target=TARGET_COLS",
         "12 целевых переменных — в энкодере как лаговые входы"],
        ["ИТОГО", str(n_enc_total), "-",
         "Полный набор признаков, доступных энкодеру за ретроспективный период"],
    ],
    col_widths=[4.5, 2, 4.5, 7.5],
    header_color="1F3964",
)

doc.add_paragraph()
h2(doc, f"6.3 Признаки декодера (будущее)")
add_table(
    doc,
    headers=["Компонент", "Кол-во", "Параметр датасета", "Описание"],
    rows=[
        ["known_categoricals", str(n_known_cats),
         "time_varying_known_categoricals",
         "season_enc, holiday_name_enc"],
        ["known_reals", str(n_known_reals),
         "time_varying_known_reals",
         f"Те же {n_known_reals} признаков, что и в энкодере — задаются на весь горизонт прогноза"],
        ["relative_time_idx", str(n_rel_time),
         "add_relative_time_idx=True",
         "Относительная позиция шага в декодерном окне [0..1]"],
        ["ИТОГО", str(n_dec_total), "-",
         "Декодер видит ТОЛЬКО то, что известно заранее — нет unknown_reals и нет целевых"],
    ],
    col_widths=[4.5, 2, 4.5, 7.5],
    header_color="1F3964",
)

doc.add_paragraph()
para(
    doc,
    "Ключевое отличие энкодера от декодера: декодер не видит unknown_reals (18 шт.) и "
    "не видит прошлых целевых значений (12 шт.) — их суммарно 30 признаков доступны "
    "только в ретроспективе. "
    f"Благодаря переносу трафика-счётчиков и цен конкурентов в KNOWN, "
    f"декодер располагает полным контекстом для сценарного прогнозирования "
    f"на {PREDICTION_LENGTH}-дневный горизонт.",
    italic=True,
)

# ============================================================
# 7. ИТОГОВАЯ СВОДКА
# ============================================================
h1(doc, "7. Итоговая сводка")

add_table(
    doc,
    headers=["TFT-роль", "Кол-во", "Примечание"],
    rows=[
        ["static_categoricals",
         str(len(STATIC_CATS)),
         "road_type, road_level, direction, settlement_size, distance_to_city_km"],
        ["static_reals",
         str(len(STATIC_REALS)),
         "shop_area_m2, num_pumps_* (10), has_* (9), competitors_wink"],
        ["time_varying_known_categoricals",
         str(len(TIME_VARYING_KNOWN_CATS)),
         "season, holiday_name"],
        ["time_varying_known_reals",
         str(len(TIME_VARYING_KNOWN_REALS)),
         "Время (9), флаги (2), weather_condition (1), акции (3), "
         "цены топлива (5), погода (2), трафик поп. (12), трафик встр. (12), конкуренты (7)"],
        ["time_varying_unknown_reals",
         str(len(TIME_VARYING_UNKNOWN_REALS)),
         "Клиентские метрики (2) + производные трафика (16)"],
        ["target (12 совместных прогнозов)",
         str(len(TARGET_COLS)),
         "5 видов топлива (AI92/95/DT/DT_bio/AI100_bio) + 7 категорий магазина"],
        ["ИТОГО входных переменных",
         str(len(STATIC_CATS) + len(STATIC_REALS) + len(TIME_VARYING_KNOWN_CATS)
             + len(TIME_VARYING_KNOWN_REALS) + len(TIME_VARYING_UNKNOWN_REALS) + len(TARGET_COLS)),
         "Без учёта time_idx, station_id и автодобавляемых target_scales/enc_length/rel_time"],
    ],
    col_widths=[5.5, 2.5, 10.5],
    header_color="2E74B5",
)

doc.add_paragraph()
add_table(
    doc,
    headers=["Тензор", "Теоретический размер", "Содержимое"],
    rows=[
        ["static",
         str(n_static_total),
         f"static_cats({n_static_cats}) + static_reals({n_static_reals}) "
         f"+ target_scales({n_target_scales}) + enc_length({n_enc_length})"],
        ["encoder_cont (прошлое)",
         str(n_enc_total),
         f"known_cats({n_known_cats}) + known_reals({n_known_reals}) "
         f"+ rel_idx({n_rel_time}) + unknown({n_unknown}) + targets_past({n_targets_enc})"],
        ["decoder_cont (будущее)",
         str(n_dec_total),
         f"known_cats({n_known_cats}) + known_reals({n_known_reals}) + rel_idx({n_rel_time})"],
        ["Реальная ширина тензора (encoder_cont = decoder_cont)",
         "118",
         "Фактическое значение из batch-проверки prepare_dataset.py. "
         "pytorch_forecasting добавляет дополнительные внутренние признаки (target_scales, "
         "enc_length broadcast per step, padding для выравнивания encoder/decoder)."],
    ],
    col_widths=[4.5, 3, 11],
    header_color="2E74B5",
)

# ── Примечание о what-if ─────────────────────────────────────
doc.add_paragraph()
para(
    doc,
    "Сценарный анализ (what-if): изменять для прогноза можно только переменные из "
    "TIME_VARYING_KNOWN_REALS и TIME_VARYING_KNOWN_CATS. "
    "Переменные из UNKNOWN_REALS (клиентские метрики, производные трафика) "
    "недоступны для декодера — модель не знает их будущих значений. "
    "Статические переменные (паспорт АЗС) зашиты в веса модели при обучении "
    "и не могут быть изменены без переобучения.",
    italic=True,
)

# ============================================================
# Сохранение
# ============================================================
out_path = "reports/tft_report.docx"
doc.save(out_path)
print(f"Отчёт сохранён: {out_path}")

print("\n" + "=" * 60)
print("Готово.")
print(f"  Сохранено : {out_path}")
print("=" * 60)
