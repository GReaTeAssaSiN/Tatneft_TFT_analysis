"""
Формирование DOCX-отчёта по TFT-анализу данных.
Описывает каждую переменную, её предобработку и роль в TFT-модели.

Запускать из корня проекта: python eda/tft_report.py
Результат: reports/tft_report.docx
"""

import os
import pickle
import sys

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from utils.data_utils import CYCLICAL_FEATURES, LOG_COLS, TARGET_COLS

os.makedirs("reports", exist_ok=True)

# ── Загрузка данных ──────────────────────────────────────────
df = pd.read_csv("data/prepared_data.csv", parse_dates=["timestamp"])
merged = pd.read_csv("data/merged_data.csv", parse_dates=["timestamp"])
train_df = pd.read_csv("data/train.csv", parse_dates=["timestamp"])
val_df = pd.read_csv("data/val.csv", parse_dates=["timestamp"])
test_df = pd.read_csv("data/test.csv", parse_dates=["timestamp"])

with open("tft/scalers.pkl", "rb") as f:
    scaler_data = pickle.load(f)
scalers = scaler_data["scalers"]
log1p_cols = scaler_data["log1p_cols"]


# ── Вспомогательные функции ──────────────────────────────────
def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x37, 0x86, 0x8C)
    return p


def add_table(doc, headers, rows, col_widths=None, header_color="1F3964"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Заголовок
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        cell_bg(hdr[i], header_color)
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

    # Данные
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cell_bg(cells[ci], bg)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(9)

    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return tbl


def para(doc, text, bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


# ============================================================
# Документ
# ============================================================
doc = Document()

# Поля
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

# ── Титул ────────────────────────────────────────────────────
title = doc.add_heading("Анализ данных и предобработка для TFT-модели", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    run.font.size = Pt(18)

subtitle = doc.add_paragraph("Прогнозирование продаж топлива сети АЗС Татнефть")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph(
    f"Temporal Fusion Transformer (Lim et al., 2020)  |  "
    f"Период: 2023 год  |  Станций: {merged['station_id'].nunique()}"
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============================================================
# 1. ИСТОЧНИКИ ДАННЫХ
# ============================================================
h1(doc, "1. Источники данных")

para(doc, "В работе используются два исходных файла для 5 АЗС:", bold=True)

add_table(
    doc,
    headers=["Файл", "Строк", "Колонок", "Содержимое"],
    rows=[
        [
            "5stations_metadata.csv",
            "5",
            "32",
            "Паспорт АЗС: тип дороги, услуги, ценовые ориентиры, кол-во колонок",
        ],
        [
            "5stations_data.csv",
            "43 800",
            "72",
            "Почасовые данные 2023 г.: погода, трафик, продажи, цены, акции",
        ],
        [
            "merged_data.csv (JOIN)",
            str(merged.shape[0]),
            str(merged.shape[1]),
            "Объединённый датасет по station_id",
        ],
        [
            "prepared_data.csv (после предобработки)",
            str(df.shape[0]),
            str(df.shape[1]),
            "Нормализованный датасет для подачи в TFT",
        ],
    ],
    col_widths=[5.5, 2, 2.5, 8],
)

doc.add_paragraph()
para(
    doc,
    "Объединение выполняется через LEFT JOIN по полю station_id. "
    "Все 43 800 строк временного ряда сохраняются; к ним добавляются "
    f"{merged.shape[1] - 72} статических признаков из паспорта АЗС.",
)

# ============================================================
# 2. КЛАССИФИКАЦИЯ ПЕРЕМЕННЫХ
# ============================================================
h1(doc, "2. Классификация переменных для TFT-модели")

para(
    doc,
    "Согласно статье Lim et al. (2020), TFT принимает три типа входов:",
    bold=True,
)

add_table(
    doc,
    headers=["Тип входа TFT", "Определение", "Кол-во колонок"],
    rows=[
        [
            "static_categoricals",
            "Категориальные свойства АЗС, не меняющиеся во времени",
            "3",
        ],
        [
            "static_reals",
            "Числовые свойства АЗС из паспорта (metadata)",
            "27",
        ],
        [
            "time_varying_known_categoricals",
            "Категориальные признаки, известные заранее (сезон, праздник)",
            "4",
        ],
        [
            "time_varying_known_reals",
            "Числовые признаки, известные заранее (цены, акции, час суток)",
            "20 + 8 sin/cos = 28",
        ],
        [
            "time_varying_unknown_reals",
            "Наблюдаемые прошлые (погода, трафик, shop_total_revenue, цены конкурентов) + TARGET_COLS",
            "19 + 12",
        ],
        [
            "target",
            "Целевые переменные: 7 видов топлива + 5 категорий магазина",
            "12",
        ],
    ],
    col_widths=[5, 10, 3.5],
)

# ── 2.1 Статические переменные ───────────────────────────────
h2(doc, "2.1 Статические переменные (не меняются во времени)")
para(
    doc,
    "Статические данные берутся из файла 5stations_metadata.csv и "
    "присоединяются к каждой строке временного ряда при JOIN.",
)

add_table(
    doc,
    headers=["Переменная", "Тип", "TFT-вход", "Предобработка"],
    rows=[
        ["road_type", "категориальная", "static_categoricals", "LabelEncoder -> road_type_enc"],
        ["direction", "категориальная", "static_categoricals", "LabelEncoder -> direction_enc"],
        ["settlement_size", "категориальная", "static_categoricals", "LabelEncoder -> settlement_size_enc"],
        ["distance_to_city_km", "вещественная", "static_reals", "Паспортные данные. Без изменений."],
        ["total_pumps", "вещественная", "static_reals", "Паспортные данные. Без изменений."],
        ["shop_area_m2", "вещественная", "static_reals", "Паспортные данные. Без изменений."],
        ["num_pumps_AI92/95/98", "вещественная", "static_reals", "Паспортные данные. Без изменений. (3 колонки)"],
        ["num_pumps_DT_*", "вещественная", "static_reals", "Паспортные данные. Без изменений. (4 колонки)"],
        ["has_car_wash/cafe/shop/...", "бинарная", "static_reals", "Паспортные данные. Бинарный 0/1."],
        ["competitors_within_5km", "вещественная", "static_reals", "Паспортные данные. Без изменений."],
        ["customer/staff/corporate scores", "вещественная", "static_reals", "Паспортные данные. Без изменений. (4 колонки)"],
        ["base_price_AI92/95/98/DT_*", "вещественная", "static_reals", "Паспортные данные. Без изменений. (7 колонок)"],
    ],
    col_widths=[5, 3, 4.5, 6],
)

# ── 2.2 Известные будущие переменные ─────────────────────────
h2(doc, "2.2 Известные будущие переменные")
para(
    doc,
    "Эти признаки можно знать заранее: расписание, цены, плановые акции. "
    "TFT подаёт их и в encoder (ретроспектива), и в decoder (прогнозный горизонт), "
    "что позволяет модели учитывать будущий контекст при построении прогноза.",
)

add_table(
    doc,
    headers=["Переменная", "Тип", "TFT-вход", "Предобработка"],
    rows=[
        ["season", "категориальная", "known_cats", "LabelEncoder -> season_enc"],
        ["day_name", "категориальная", "known_cats", "LabelEncoder -> day_name_enc"],
        ["ad_channel", "категориальная", "known_cats", "NaN -> 'нет_рекламы'. LabelEncoder"],
        ["holiday_name", "категориальная", "known_cats", "NaN -> 'нет_праздника'. LabelEncoder"],
        ["hour", "вещественная (цикл.)", "known_reals", "Z-score + hour_sin, hour_cos (period=24)"],
        ["day_of_week", "вещественная (цикл.)", "known_reals", "Z-score + dow_sin, dow_cos (period=7)"],
        ["month", "вещественная (цикл.)", "known_reals", "Z-score + month_sin, month_cos (period=12)"],
        ["week_of_year", "вещественная (цикл.)", "known_reals", "Z-score + woy_sin, woy_cos (period=52)"],
        ["quarter", "вещественная", "known_reals", "Z-score"],
        ["is_weekend / is_holiday / ...", "бинарная", "known_reals", "Без нормализации (0/1)"],
        ["promotion_fuel/shop/cafe_active", "бинарная", "known_reals", "Без нормализации (0/1)"],
        ["ad_active", "бинарная", "known_reals", "Без нормализации (0/1)"],
        ["price_AI92/95/98/DT_*", "вещественная", "known_reals", "Z-score per-station (7 колонок)"],
    ],
    col_widths=[5.5, 3.5, 3, 6.5],
)

# ── 2.3 Наблюдаемые прошлые ──────────────────────────────────
h2(doc, "2.3 Наблюдаемые прошлые (observed past inputs)")
para(
    doc,
    "Эти признаки доступны только за прошлое — их нельзя знать заранее. "
    "TFT использует их только в encoder (168 часов ретроспективы). "
    "Включают погоду, дорожный трафик, выручку магазина и цены конкурентов.",
)

add_table(
    doc,
    headers=["Группа", "Переменные", "Предобработка"],
    rows=[
        [
            "Погода (8 переменных)",
            "temperature, precipitation_mm, visibility_km, wind_speed_ms,\nis_snow, is_rain, is_fog, weather_condition",
            "Winsorization IQR + Z-score. weather_condition -> float через _enc",
        ],
        [
            "Трафик (7 переменных)",
            "traffic_Passengers_cars, traffic_Truck_short, traffic_Truck,\ntraffic_Truck_long, traffic_Transporter, traffic_Undefined, total_traffic",
            "Winsorization IQR + Z-score",
        ],
        [
            "Магазин (1 переменная)",
            "shop_total_revenue",
            "log1p + _orig. Z-score не применяется (в LOG_COLS — нормализует TorchNormalizer внутри TFT).\n"
            "Категории магазина (shop_напитки, shop_закуски, shop_автотовары,\n"
            "shop_кофе, shop_табак) перенесены в целевые переменные.",
        ],
        [
            "Цены конкурентов (3 переменных)",
            "competitor_price_AI92, competitor_price_AI95, competitor_price_DT",
            "Winsorization IQR + Z-score",
        ],
    ],
    col_widths=[4.5, 8, 6],
)

# ── 2.4 Целевые переменные ────────────────────────────────────
h2(doc, "2.4 Целевые переменные (выходы модели)")
para(
    doc,
    "12 переменных являются одновременно целью прогноза (выход модели) "
    "и наблюдаемыми прошлыми входами (encoder видит их лаговые значения). "
    "Это стандартная авторегрессионная схема TFT (Section 3, Lim et al., 2020):\n"
    "  - 7 видов топлива: sales_AI92/95/98, sales_DT_EURO/TANEKO/SUMMER/WINTER\n"
    "  - 5 категорий магазина: shop_напитки, shop_закуски, shop_автотовары, shop_кофе, shop_табак\n"
    "Модель прогнозирует продажи топлива и сопутствующих товаров совместно, "
    "используя взаимные корреляции между группами (покупатели топлива = покупатели магазина).",
)

# Вычисляем skew до/после
skew_rows = []
for col in TARGET_COLS:
    orig_col = col + "_orig"
    if orig_col in df.columns:
        skew_before = df[orig_col].skew()
        skew_after = df[col].skew()
        mn = df[orig_col].min()
        mx = df[orig_col].max()
        skew_rows.append([col, f"{mn:.0f} – {mx:.0f}", f"{skew_before:.2f}", f"{skew_after:.2f}"])
    else:
        skew_rows.append([col, "—", "—", "—"])

add_table(
    doc,
    headers=["Переменная", "Диапазон (л/ч)", "Skew до log1p", "Skew после log1p"],
    rows=skew_rows,
    col_widths=[5, 4, 4, 4],
)

doc.add_paragraph()
para(
    doc,
    "Дополнительная нормализация целевых переменных выполняется TorchNormalizer(method='robust') "
    "внутри TFT — этот нормализатор использует медиану и IQR вместо mean/std, "
    "что делает его устойчивым к оставшимся выбросам.",
)

# ============================================================
# 3. ПРЕДОБРАБОТКА: ПОШАГОВОЕ ОПИСАНИЕ
# ============================================================
h1(doc, "3. Предобработка данных")

steps = [
    (
        "Шаг 1. Заполнение пропусков",
        "Два поля содержат пропуски: holiday_name (42 720 из 43 800 строк) "
        "и ad_channel (30 609 строк). Отсутствие значения несёт смысл:\n"
        "  - holiday_name -> 'нет_праздника' (обычный рабочий день)\n"
        "  - ad_channel -> 'нет_рекламы' (реклама не размещалась)\n"
        "Удаление строк с NaN недопустимо — нарушает непрерывность временного ряда.",
    ),
    (
        "Шаг 2. Устранение выбросов (Winsorization)",
        "Метод IQR: значения, выходящие за [Q1 - 1.5*IQR, Q3 + 1.5*IQR], "
        "заменяются граничными значениями. Строки не удаляются.\n"
        "Применяется к интервальным (не бинарным) переменным временного ряда.\n"
        "Исключены: бинарные (0/1), категориальные, а также все статические "
        "переменные из metadata — они являются паспортными характеристиками АЗС "
        "и не могут быть выбросами.",
    ),
    (
        "Шаг 3. Label Encoding категориальных переменных",
        "Категориальные строковые переменные кодируются целыми числами (LabelEncoder).\n"
        "TFT строит Embedding-слои по целым числам — это эффективнее One-Hot Encoding.\n"
        "Создаются новые колонки с суффиксом _enc. Исходные строковые колонки сохранены.",
    ),
    (
        "Шаг 4. Циклическое sin/cos кодирование",
        "Временные признаки hour, day_of_week, month, week_of_year имеют циклическую природу:\n"
        "  - hour=23 и hour=0 — соседние часы, но числово далеки (разница 23)\n"
        "  - Решение: sin/cos на единичной окружности\n"
        "    hour_sin = sin(2π * hour / 24)\n"
        "    hour_cos = cos(2π * hour / 24)\n"
        "  => ||(23_sin, 23_cos) - (0_sin, 0_cos)||₂ минимален\n"
        "Добавляются 8 новых колонок (2 * 4 признака).",
    ),
    (
        "Шаг 5. Монотонный индекс времени (time_idx)",
        "TFT требует целочисленный, возрастающий time_idx без пропусков для каждой группы.\n"
        "Реализуется через cumcount() после сортировки по (station_id, timestamp):\n"
        "  time_idx = 0, 1, 2, ..., 8759 — для каждой из 5 станций",
    ),
    (
        "Шаг 6. Логарифмическое преобразование (log1p)",
        "Все целевые переменные (7 видов топлива + 5 категорий магазина) и shop_total_revenue "
        "имеют правостороннюю асимметрию (skew > 1). log1p(x) = log(1+x):\n"
        "  - работает при x=0 (ночные часы — продажи равны нулю)\n"
        "  - снижает skew и стабилизирует дисперсию\n"
        "  - улучшает сходимость нейросети\n"
        "Оригинальные значения сохраняются с суффиксом _orig для обратного преобразования.",
    ),
    (
        "Шаг 7. Z-score нормализация per-station",
        "Каждая станция нормализуется независимо: (x - mean_s) / std_s.\n"
        "Обоснование: трассовые АЗС имеют масштаб продаж в 3–5 раз выше городских.\n"
        "Исключены из нормализации:\n"
        "  - статические переменные из metadata (паспортные данные АЗС не меняются)\n"
        "  - бинарные переменные (0/1)\n"
        "  - _enc колонки (целые числа для embedding)\n"
        "  - _orig колонки (оригинальные значения до log1p)\n"
        "  - LOG_COLS (целевые переменные — нормализует TorchNormalizer внутри TFT)\n"
        "Скейлеры (mean, std) сохраняются в tft/scalers.pkl для обратного преобразования.",
    ),
    (
        "Шаг 8. Темпоральный сплит",
        "Данные разбиваются строго по времени (без рандомизации):\n"
        f"  Train : Jan – Oct 2023  ({len(train_df)} строк, {len(train_df)/len(df)*100:.1f}%)\n"
        f"  Val   : Nov 2023        ({len(val_df)} строк, {len(val_df)/len(df)*100:.1f}%)\n"
        f"  Test  : Dec 2023        ({len(test_df)} строк, {len(test_df)/len(df)*100:.1f}%)\n"
        "Случайная выборка недопустима — вызвала бы data leakage из будущего в прошлое.",
    ),
]

for title_step, desc in steps:
    h3(doc, title_step)
    for line in desc.split("\n"):
        p = doc.add_paragraph(line, style="List Bullet" if line.startswith("  -") else "Normal")
        for run in p.runs:
            run.font.size = Pt(10)

# ============================================================
# 4. ВХОДЫ И ВЫХОДЫ TFT
# ============================================================
h1(doc, "4. Итоговая схема входов и выходов TFT")

h2(doc, "4.1 Параметры модели")

add_table(
    doc,
    headers=["Параметр", "Значение", "Описание"],
    rows=[
        ["ENCODER_LENGTH", "168 ч", "Глубина ретроспективы (7 суток)"],
        ["PREDICTION_LENGTH", "24 ч", "Горизонт прогноза (1 сутки)"],
        ["hidden_size", "64 / 128", "CPU / GPU"],
        ["attention_head_size", "2 / 4", "CPU / GPU"],
        ["hidden_continuous_size", "64", "Размер сети для числовых входов"],
        ["dropout", "0.15", "Регуляризация"],
        ["learning_rate", "3e-4", "Начальный LR (ReduceOnPlateau, patience=5)"],
        ["gradient_clip", "1.0", "Обрезка градиентов"],
        ["EarlyStopping", "patience=12", "Останавливается при отсутствии улучшения"],
        ["target_normalizer", "TorchNormalizer(robust)", "Медиана + IQR, устойчив к выбросам"],
    ],
    col_widths=[5, 4, 9.5],
)

h2(doc, "4.2 Входные данные")

add_table(
    doc,
    headers=["TFT-вход", "Источник", "Колонки (после предобработки)"],
    rows=[
        [
            "group_ids",
            "5stations_data.csv",
            "station_id (строка)",
        ],
        [
            "static_categoricals",
            "metadata",
            "road_type_enc, direction_enc, settlement_size_enc",
        ],
        [
            "static_reals",
            "metadata",
            "distance_to_city_km, total_pumps, shop_area_m2, "
            "num_pumps_* (7), has_* (5), scores (4), base_price_* (7) — всего 27",
        ],
        [
            "known_categoricals",
            "data",
            "season_enc, day_name_enc, ad_channel_enc, holiday_name_enc",
        ],
        [
            "known_reals",
            "data",
            "hour (+sin/cos), day_of_week (+sin/cos), week_of_year (+sin/cos), "
            "month (+sin/cos), quarter, is_* (4), promotion_* (3), "
            "ad_active, price_* (7) — итого 28 колонок",
        ],
        [
            "unknown_reals",
            "data",
            "погода (8), трафик (7), магазин (1: shop_total_revenue), competitor_price (3) + TARGET_COLS (12) — итого 31",
        ],
    ],
    col_widths=[3.5, 3, 12],
)

h2(doc, "4.3 Выходные данные")
para(
    doc,
    "TFT выдаёт квантильные прогнозы (QuantileLoss) для каждой из 12 целевых переменных "
    "на 24 часа вперёд (7 видов топлива + 5 категорий магазина). "
    "Квантили по умолчанию: q2, q10, q25, q50, q75, q90, q98 — позволяют строить "
    "доверительные интервалы прогноза и оценивать неопределённость.",
)

fuel_cols = [c for c in TARGET_COLS if c.startswith("sales_")]
shop_cols = [c for c in TARGET_COLS if c.startswith("shop_")]

add_table(
    doc,
    headers=["Переменная", "Группа", "Единица", "Описание"],
    rows=(
        [[col, "Топливо", "л/ч", "Продажи топлива за 1 час"] for col in fuel_cols]
        + [[col, "Магазин", "руб/ч", "Выручка категории магазина за 1 час"] for col in shop_cols]
    ),
    col_widths=[5, 3, 2.5, 8],
)

para(
    doc,
    "\nОбратное преобразование прогноза (tft/predict.py): "
    "TorchNormalizer^-1 (mode='quantiles', автоматически) -> expm1(x) -> исходные единицы (литры/час для топлива, руб/ч для магазина).",
)

# ============================================================
# 5. ИТОГОВЫЙ СОСТАВ ПЕРЕМЕННЫХ МОДЕЛИ
# ============================================================
h1(doc, "5. Итоговый состав переменных модели после предобработки")

para(
    doc,
    "Таблица показывает, какие именно колонки подаются в TFT после всех шагов "
    "предобработки. Указаны применённые преобразования и роль в модели.",
)

VAR_TABLE = [
    # (Колонка в модели, Исходная переменная, Преобразование, TFT-роль)
    # static_categoricals
    ("road_type_enc",               "road_type",               "LabelEncoder",                        "static_cat"),
    ("direction_enc",               "direction",               "LabelEncoder",                        "static_cat"),
    ("settlement_size_enc",         "settlement_size",         "LabelEncoder",                        "static_cat"),
    # static_reals
    ("distance_to_city_km",         "distance_to_city_km",     "—",                                   "static_real"),
    ("total_pumps",                 "total_pumps",             "—",                                   "static_real"),
    ("shop_area_m2",                "shop_area_m2",            "—",                                   "static_real"),
    ("num_pumps_AI92/95/98 (3 к.)", "num_pumps_AI92/95/98",   "— (3 колонки)",                       "static_real"),
    ("num_pumps_DT_* (4 к.)",       "num_pumps_DT_*",          "— (4 колонки)",                       "static_real"),
    ("has_* (5 к.)",                "has_car_wash/cafe/...",   "— (бинарные)",                        "static_real"),
    ("competitors_within_5km",      "competitors_within_5km",  "—",                                   "static_real"),
    ("*_score (4 к.)",              "customer/staff/corp/svc", "— (4 колонки)",                       "static_real"),
    ("base_price_* (7 к.)",         "base_price_AI92/...",     "— (7 колонок)",                       "static_real"),
    # known_categoricals
    ("season_enc",                  "season",                  "LabelEncoder",                        "known_cat"),
    ("day_name_enc",                "day_name",                "LabelEncoder",                        "known_cat"),
    ("ad_channel_enc",              "ad_channel",              "NaN→'нет_рекламы' + LabelEncoder",    "known_cat"),
    ("holiday_name_enc",            "holiday_name",            "NaN→'нет_праздника' + LabelEncoder",  "known_cat"),
    # known_reals
    ("hour",                        "hour",                    "Z-score",                             "known_real"),
    ("hour_sin / hour_cos",         "hour",                    "sin/cos(2π·h/24)",                    "known_real"),
    ("day_of_week",                 "day_of_week",             "Z-score",                             "known_real"),
    ("dow_sin / dow_cos",           "day_of_week",             "sin/cos(2π·d/7)",                     "known_real"),
    ("month",                       "month",                   "Z-score",                             "known_real"),
    ("month_sin / month_cos",       "month",                   "sin/cos(2π·m/12)",                    "known_real"),
    ("week_of_year",                "week_of_year",            "Z-score",                             "known_real"),
    ("woy_sin / woy_cos",           "week_of_year",            "sin/cos(2π·w/52)",                    "known_real"),
    ("quarter",                     "quarter",                 "Z-score",                             "known_real"),
    ("is_weekend",                  "is_weekend",              "— (бинарная)",                        "known_real"),
    ("is_holiday",                  "is_holiday",              "— (бинарная)",                        "known_real"),
    ("is_rush_hour",                "is_rush_hour",            "— (бинарная)",                        "known_real"),
    ("is_night",                    "is_night",                "— (бинарная)",                        "known_real"),
    ("promotion_fuel_active",       "promotion_fuel_active",   "— (бинарная)",                        "known_real"),
    ("promotion_shop_active",       "promotion_shop_active",   "— (бинарная)",                        "known_real"),
    ("promotion_cafe_active",       "promotion_cafe_active",   "— (бинарная)",                        "known_real"),
    ("ad_active",                   "ad_active",               "— (бинарная)",                        "known_real"),
    ("price_AI92/95/98 (3 к.)",     "price_AI92/95/98",        "Z-score per-station",                 "known_real"),
    ("price_DT_* (4 к.)",           "price_DT_*",              "Z-score per-station (4 кол.)",         "known_real"),
    # unknown_reals (observed past)
    ("temperature",                 "temperature",             "Winsorization + Z-score",             "unknown_real"),
    ("precipitation_mm",            "precipitation_mm",        "Winsorization + Z-score",             "unknown_real"),
    ("visibility_km",               "visibility_km",           "Winsorization + Z-score",             "unknown_real"),
    ("wind_speed_ms",               "wind_speed_ms",           "Winsorization + Z-score",             "unknown_real"),
    ("is_snow / is_rain / is_fog",  "is_snow/rain/fog",        "— (бинарные, 3 кол.)",                "unknown_real"),
    ("weather_condition",           "weather_condition",       "Winsorization + Z-score",             "unknown_real"),
    ("traffic_Passengers_cars",     "traffic_Passengers_cars", "Winsorization + Z-score",             "unknown_real"),
    ("traffic_Truck_short",         "traffic_Truck_short",     "Winsorization + Z-score",             "unknown_real"),
    ("traffic_Truck",               "traffic_Truck",           "Winsorization + Z-score",             "unknown_real"),
    ("traffic_Truck_long",          "traffic_Truck_long",      "Winsorization + Z-score",             "unknown_real"),
    ("traffic_Transporter",         "traffic_Transporter",     "Winsorization + Z-score",             "unknown_real"),
    ("traffic_Undefined",           "traffic_Undefined",       "Winsorization + Z-score",             "unknown_real"),
    ("total_traffic",               "total_traffic",           "Winsorization + Z-score",             "unknown_real"),
    ("shop_total_revenue",          "shop_total_revenue",      "log1p",                               "unknown_real"),
    ("competitor_price_AI92",       "competitor_price_AI92",   "Winsorization + Z-score",             "unknown_real"),
    ("competitor_price_AI95",       "competitor_price_AI95",   "Winsorization + Z-score",             "unknown_real"),
    ("competitor_price_DT",         "competitor_price_DT",     "Winsorization + Z-score",             "unknown_real"),
    # targets (также observed past в encoder)
    ("sales_AI92/95/98 (3 к.)",     "sales_AI92/95/98",        "log1p + TorchNormalizer(robust)",     "target"),
    ("sales_DT_* (4 к.)",           "sales_DT_EURO/...",       "log1p + TorchNormalizer(robust)",     "target"),
    ("shop_* (5 к.)",               "shop_напитки/...",         "log1p + TorchNormalizer(robust)",     "target"),
]

ROLE_LABELS = {
    "static_cat":   "static_categoricals",
    "static_real":  "static_reals",
    "known_cat":    "known_categoricals",
    "known_real":   "known_reals",
    "unknown_real": "observed past (unknown_reals)",
    "target":       "target + observed past в encoder",
}

add_table(
    doc,
    headers=["Колонка в модели", "Исходная переменная", "Преобразование", "TFT-роль"],
    rows=[
        [col, src, transform, ROLE_LABELS[role]]
        for col, src, transform, role in VAR_TABLE
    ],
    col_widths=[4.5, 4.0, 4.5, 5.5],
)

doc.add_paragraph()
h2(doc, "5.1 Итоговое количество переменных по группам")

add_table(
    doc,
    headers=["TFT-роль", "Кол-во колонок", "Примечание"],
    rows=[
        ["static_categoricals",             "3",   "Тип дороги, направление, размер поселения"],
        ["static_reals",                    "27",  "Паспортные характеристики АЗС"],
        ["known_categoricals",              "4",   "Сезон, день недели, канал рекламы, праздник"],
        ["known_reals",                     "28",  "Время (8 цикл.), бинарные (7), цены (7), прочие (6)"],
        ["observed past (unknown_reals)",   "19",  "Погода (8), трафик (7), магазин (1), конкуренты (3)"],
        ["target / observed past encoder",  "12",  "7 видов топлива + 5 категорий магазина"],
        ["ИТОГО входных колонок",           str(3+27+4+28+19+12), "Без учёта time_idx и station_id"],
    ],
    col_widths=[5.5, 3.0, 10.0],
    header_color="2E74B5",
)


doc.add_paragraph()
h2(doc, "5.2 Из чего складываются 55 статических, 64 прошлых, 33 будущих")

para(
    doc,
    "Числа 55 / 64 / 33 — это фактические размеры тензоров, которые TFT получает "
    "на вход. Они включают не только исходные переменные, но и дополнительные "
    "признаки, автоматически добавляемые pytorch-forecasting при построении датасета.",
)

doc.add_paragraph()
h3(doc, "Статические признаки — 55")
add_table(
    doc,
    headers=["Компонент", "Кол-во", "Параметр датасета", "Описание"],
    rows=[
        ["static_categoricals",      "3",   "—",                           "road_type_enc, direction_enc, settlement_size_enc"],
        ["static_reals (паспорт АЗС)", "27", "—",                          "Расстояния, колонки, площадь, сервисы, базовые цены, оценки"],
        ["target scales",            "24",  "add_target_scales=True",      "12 целей × 2: location (центр) + scale (масштаб) у TorchNormalizer(robust)"],
        ["encoder_length",           "1",   "add_encoder_length=True",     "Фактическая длина энкодера в текущем сэмпле — подсказка для модели"],
        ["ИТОГО",                    "55",  "—",                           ""],
    ],
    col_widths=[4.5, 2.0, 4.5, 7.5],
    header_color="1F3964",
)

doc.add_paragraph()
h3(doc, "Прошлые признаки (encoder_cont) — 64")
add_table(
    doc,
    headers=["Компонент", "Кол-во", "Параметр датасета", "Описание"],
    rows=[
        ["known_categoricals",       "4",   "—",                           "season_enc, day_name_enc, ad_channel_enc, holiday_name_enc"],
        ["known_reals",              "28",  "—",                           "Время (raw+sin/cos ×4), бинарные флаги (4), акции (4), цены топлива (7)"],
        ["relative_time_idx",        "1",   "add_relative_time_idx=True",  "Позиция шага внутри окна (0 = начало, 1 = конец)"],
        ["unknown_reals",            "19",  "—",                           "Погода (8), трафик (7), выручка магазина (1), цены конкурентов (3)"],
        ["targets (observed past)",  "12",  "—",                           "12 целевых переменных — в энкодере видны как прошлые значения"],
        ["ИТОГО",                    "64",  "—",                           "= 33 (известные) + 19 (наблюдаемые) + 12 (целевые прошлые)"],
    ],
    col_widths=[4.5, 2.0, 4.5, 7.5],
    header_color="1F3964",
)

doc.add_paragraph()
h3(doc, "Будущие признаки (decoder_cont) — 33")
add_table(
    doc,
    headers=["Компонент", "Кол-во", "Параметр датасета", "Описание"],
    rows=[
        ["known_categoricals",  "4",   "—",                           "season_enc, day_name_enc, ad_channel_enc, holiday_name_enc"],
        ["known_reals",         "28",  "—",                           "Время (raw+sin/cos ×4), бинарные флаги (4), акции (4), цены топлива (7)"],
        ["relative_time_idx",   "1",   "add_relative_time_idx=True",  "Позиция шага внутри горизонта прогноза"],
        ["ИТОГО",               "33",  "—",                           "Декодер видит только то, что можно знать заранее"],
    ],
    col_widths=[4.5, 2.0, 4.5, 7.5],
    header_color="1F3964",
)

para(
    doc,
    "Ключевое отличие: декодер (33) — это подмножество энкодера (64). "
    "В будущем нельзя наблюдать погоду, трафик и фактические продажи, "
    "поэтому unknown_reals и targets в decoder_cont отсутствуют. "
    "Именно это делает TFT мощнее классических seq2seq: "
    "модель явно разделяет, что ей известно заранее, а что — только по истории.",
)

# ============================================================
# Сохранение
# ============================================================
out_path = "reports/tft_report.docx"
doc.save(out_path)
print(f"Отчёт сохранён: {out_path}")

print("\n" + "=" * 60)
print("Готово.")
print("  Сохранено : reports/tft_report.docx")
print("  Следующий шаг: python tft/prepare_dataset.py")
print("=" * 60)
